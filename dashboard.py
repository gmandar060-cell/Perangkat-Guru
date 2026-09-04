import streamlit as st
from google import genai
from datetime import datetime
import re
import io
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# =========================================================
# SIAP AJAR 22 - DASHBOARD
# =========================================================

if not st.session_state.get("authenticated", False):
    st.stop()


# =========================================================
# SESSION STATE
# =========================================================

if "user_name" not in st.session_state:
    st.session_state.user_name = ""

if "user_api_key" not in st.session_state:
    st.session_state.user_api_key = ""

if "hasil_teks" not in st.session_state:
    st.session_state.hasil_teks = ""

if "nama_file_base" not in st.session_state:
    st.session_state.nama_file_base = "Perangkat_Kurikulum"


# =========================================================
# LOGO
# =========================================================

LOGO_URL = (
    "https://raw.githubusercontent.com/"
    "gmandar060-cell/Perangkat-Guru/main/logo.png"
)


# =========================================================
# CSS
# =========================================================

st.markdown(
    """
    <style>
    @import url(
        'https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:'
        'wght@400;500;600;700;800&display=swap'
    );

    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', sans-serif !important;
    }

    .stApp {
        background: #F8FAFC;
    }

    .block-container {
        max-width: 1180px;
        padding-top: 2rem;
        padding-bottom: 3rem;
    }

    input, textarea {
        color: #0F172A !important;
        background: #FFFFFF !important;
        -webkit-text-fill-color: #0F172A !important;
    }

    input::placeholder,
    textarea::placeholder {
        color: #94A3B8 !important;
        opacity: 1 !important;
    }

    label {
        color: #1E293B !important;
        font-weight: 600 !important;
    }

    .brand-logo {
        text-align: center;
        margin-bottom: 8px;
    }

    .brand-logo img {
        width: 70px;
        height: 70px;
        object-fit: contain;
    }

    .info-card {
        background: white;
        border: 1px solid #E2E8F0;
        border-radius: 14px;
        padding: 18px;
        margin-bottom: 15px;
    }

    .info-number {
        font-size: 25px;
        font-weight: 800;
        color: #1E3A8A !important;
    }

    .info-label {
        color: #64748B !important;
        font-size: 12px;
    }

    .section-title {
        font-size: 18px;
        font-weight: 800;
        color: #0F172A !important;
        margin-bottom: 10px;
    }

    .paper-a4 {
        background: #FFFFFF !important;
        border: 1px solid #CBD5E1;
        border-radius: 5px;
        padding: 45px 50px;
        margin: 18px auto;
        max-width: 900px;
        box-shadow: 0 12px 30px rgba(0, 0, 0, 0.08);
        color: #0F172A !important;
    }

    .paper-a4 h1,
    .paper-a4 h2,
    .paper-a4 h3,
    .paper-a4 p,
    .paper-a4 td,
    .paper-a4 th {
        color: #0F172A !important;
    }

    .paper-a4 table {
        width: 100%;
        border-collapse: collapse;
        margin: 15px 0;
        font-size: 13px;
    }

    .paper-a4 th {
        background: #E2E8F0 !important;
        border: 1px solid #64748B;
        padding: 8px;
    }

    .paper-a4 td {
        border: 1px solid #94A3B8;
        padding: 8px;
    }

    .footer-box {
        text-align: center;
        color: #64748B !important;
        font-size: 12px;
        border-top: 1px solid #E2E8F0;
        padding-top: 20px;
        margin-top: 35px;
    }

    .stButton > button,
    .stDownloadButton > button {
        border-radius: 10px !important;
        font-weight: 700 !important;
        min-height: 44px;
    }

    section[data-testid="stSidebar"] {
        background: #FFFFFF !important;
        border-right: 1px solid #E2E8F0;
    }

    @media (max-width: 768px) {
        .paper-a4 {
            padding: 22px 16px;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# FUNGSI DOCX
# =========================================================

def buat_file_docx(teks):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = teks.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []

            while i < len(lines):
                current = lines[i].strip()

                if not (
                    current.startswith("|")
                    and current.endswith("|")
                ):
                    break

                if not re.match(
                    r"^\|[\s\-:|]+\|$",
                    current
                ):
                    cells = [
                        cell.strip()
                        for cell in current[1:-1].split("|")
                    ]
                    table_lines.append(cells)

                i += 1

            if table_lines:
                rows = len(table_lines)
                cols = max(len(row) for row in table_lines)

                table = doc.add_table(
                    rows=rows,
                    cols=cols
                )

                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r, row_data in enumerate(table_lines):
                    for c in range(cols):
                        value = ""

                        if c < len(row_data):
                            value = row_data[c]

                        value = (
                            value
                            .replace("<br>", "\n")
                            .replace("<br/>", "\n")
                            .replace("**", "")
                        )

                        cell = table.cell(r, c)
                        cell.text = value

                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Calibri"
                                run.font.size = Pt(9.5)

                                if r == 0:
                                    run.font.bold = True

                doc.add_paragraph()

            continue

        if line.startswith("# "):
            paragraph = doc.add_heading(
                line[2:],
                level=1
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith("## "):
            doc.add_heading(
                line[3:],
                level=2
            )

        elif line.startswith("### "):
            doc.add_heading(
                line[4:],
                level=3
            )

        elif line.startswith("---"):
            doc.add_paragraph()

        else:
            clean = re.sub(r"<.*?>", "", line)
            clean = re.sub(
                r"\*\*(.*?)\*\*",
                r"\1",
                clean
            )

            paragraph = doc.add_paragraph(clean)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(4)

            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        i += 1

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output


# =========================================================
# MARKDOWN → HTML
# =========================================================

def markdown_to_html(teks):
    lines = teks.splitlines()
    html = []

    table_rows = []
    in_table = False

    def flush_table():
        nonlocal table_rows
        nonlocal in_table

        if not table_rows:
            return

        html.append("<table>")

        for row_index, row in enumerate(table_rows):
            tag = "th" if row_index == 0 else "td"

            html.append("<tr>")

            for cell in row:
                safe = (
                    cell
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )

                safe = re.sub(
                    r"\*\*(.*?)\*\*",
                    r"<strong>\1</strong>",
                    safe
                )

                html.append(
                    f"<{tag}>{safe}</{tag}>"
                )

            html.append("</tr>")

        html.append("</table>")

        table_rows = []
        in_table = False

    for raw in lines:
        line = raw.strip()

        if line.startswith("```"):
            continue

        if line.startswith("|") and line.endswith("|"):
            if re.match(
                r"^\|[\s\-:|]+\|$",
                line
            ):
                continue

            in_table = True

            cells = [
                cell.strip()
                for cell in line[1:-1].split("|")
            ]

            table_rows.append(cells)
            continue

        if in_table:
            flush_table()

        if not line:
            html.append("<p>&nbsp;</p>")
            continue

        escaped = (
            line
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        escaped = re.sub(
            r"\*\*(.*?)\*\*",
            r"<strong>\1</strong>",
            escaped
        )

        if escaped.startswith("### "):
            html.append(
                f"<h3>{escaped[4:]}</h3>"
            )

        elif escaped.startswith("## "):
            html.append(
                f"<h2>{escaped[3:]}</h2>"
            )

        elif escaped.startswith("# "):
            html.append(
                f"<h1>{escaped[2:]}</h1>"
            )

        elif escaped.startswith("---"):
            html.append("<hr>")

        else:
            html.append(
                f"<p>{escaped}</p>"
            )

    if in_table:
        flush_table()

    return "\n".join(html)

# =========================================================
# ATURAN 22 PERANGKAT
# =========================================================

ATURAN_22 = {

    "01. Capaian Pembelajaran (CP)": """
FORMAT WAJIB CP:
1. Identitas mata pelajaran.
2. Fase dan kelas.
3. Rasional mata pelajaran.
4. Tujuan mata pelajaran.
5. Karakteristik mata pelajaran.
6. Capaian Pembelajaran.
7. Capaian per elemen sesuai fase.
8. Catatan sumber/acuan.

ATURAN:
- CP resmi pemerintah tidak boleh dikarang.
- Jika teks CP resmi tidak tersedia dalam data pengguna,
  jangan mengklaim teks tersebut sebagai kutipan resmi.
- Bedakan antara kutipan CP resmi dan penjelasan/struktur
  pendukung.
""",

    "02. Alur Tujuan Pembelajaran (ATP)": """
FORMAT WAJIB ATP:
Tabel dengan kolom:
| No | Elemen | Capaian Pembelajaran | Tujuan Pembelajaran (TP) | Alokasi Waktu | Sumber |

ATURAN:
- TP harus diturunkan secara logis dari CP.
- Urutan TP harus sistematis dari sederhana ke kompleks.
- Alokasi waktu harus realistis.
- Jangan membuat TP yang tidak berhubungan dengan materi.
""",

    "03. Modul Ajar (RPP Plus)": """
FORMAT WAJIB MODUL AJAR:

A. INFORMASI UMUM
- Identitas modul
- Kompetensi awal
- Profil Pelajar Pancasila
- Sarana dan prasarana
- Target peserta didik
- Model/metode pembelajaran

B. KOMPONEN INTI
- Tujuan pembelajaran
- Pemahaman bermakna
- Pertanyaan pemantik
- Persiapan pembelajaran
- Langkah-langkah pembelajaran
- Kegiatan pendahuluan
- Kegiatan inti
- Kegiatan penutup
- Asesmen diagnostik
- Asesmen formatif
- Asesmen sumatif
- Diferensiasi pembelajaran
- Remedial
- Pengayaan
- Refleksi guru
- Refleksi peserta didik

C. LAMPIRAN
- LKPD
- Bahan bacaan
- Instrumen asesmen
- Rubrik penilaian
- Kunci jawaban jika diperlukan
- Glosarium
- Daftar pustaka
""",

    "04. Kriteria Ketercapaian Tujuan Pembelajaran (KKTP)": """
FORMAT WAJIB KKTP:

Tabel:
| No | Tujuan Pembelajaran | Indikator/Kriteria Ketercapaian | Bukti Asesmen | Kategori/Interval | Tindak Lanjut |

Gunakan salah satu pendekatan yang jelas:
- Rubrik
- Deskripsi kriteria
- Interval nilai

ATURAN:
- Kriteria harus dapat diamati dan diukur.
- Jelaskan indikator ketercapaian.
- Hubungkan KKTP dengan TP.
""",

    "05. Modul Projek (P5)": """
FORMAT WAJIB MODUL PROJEK:

1. Profil Modul
2. Tema Projek
3. Topik Projek
4. Dimensi Profil Pelajar Pancasila
5. Elemen dan subelemen
6. Tujuan projek
7. Target peserta didik
8. Sarana dan prasarana
9. Tahapan projek
10. Jadwal kegiatan
11. Aktivitas peserta didik
12. Peran guru
13. Asesmen diagnostik
14. Asesmen formatif
15. Asesmen sumatif
16. Rubrik penilaian
17. Refleksi
18. Tindak lanjut
""",

    "06. Kalender Pendidikan": """
FORMAT WAJIB:
Tabel kalender pendidikan yang memuat:
- Bulan
- Minggu
- Hari/tanggal
- Kegiatan
- Hari efektif
- Hari libur

ATURAN PENTING:
- Jangan mengarang tanggal resmi hari libur.
- Jangan mengklaim kalender sebagai kalender resmi
  Dinas/Sekolah jika data resmi tidak diberikan.
- Jika data tanggal resmi belum tersedia, buat FORMAT
  KALENDER SIAP DIISI dan tandai bagian yang membutuhkan
  data resmi.
""",

    "07. Program Tahunan (Prota)": """
FORMAT WAJIB PROTA:

Tabel:
| Semester | No | Tujuan Pembelajaran / Lingkup Materi | Alokasi Jam Pelajaran (JP) |

ATURAN:
- Distribusi harus mencakup satu tahun pembelajaran.
- Total JP harus konsisten dengan alokasi waktu.
- Materi harus sesuai fase, kelas dan mata pelajaran.
""",

    "08. Program Semester (Prosem)": """
FORMAT WAJIB PROSEM:

Tabel/matriks:
| No | Materi | JP | Bulan | Minggu ke-1 | Minggu ke-2 | Minggu ke-3 | Minggu ke-4 |

ATURAN:
- Distribusi JP harus realistis.
- Sesuaikan dengan semester yang dipilih.
- Jangan membuat tanggal atau pekan efektif resmi
  tanpa sumber kalender pendidikan.
""",

    "09. Jadwal Mengajar": """
FORMAT WAJIB:
Tabel:
| Hari | Jam Ke | Waktu | Kelas | Mata Pelajaran | Materi/Kegiatan | JP |

ATURAN:
- Jadwal harus mudah dibaca.
- Jangan membuat jadwal bentrok.
- Jika data jam sekolah belum tersedia, gunakan format
  jadwal yang dapat disesuaikan.
""",

    "10. Bahan Ajar": """
FORMAT WAJIB:
1. Identitas bahan ajar
2. Tujuan pembelajaran
3. Capaian pembelajaran terkait
4. Petunjuk penggunaan
5. Materi pembelajaran
6. Contoh
7. Aktivitas belajar
8. Rangkuman
9. Latihan
10. Evaluasi
11. Kunci jawaban jika diperlukan
12. Daftar pustaka

ATURAN:
- Materi harus sesuai kelas, fase dan mata pelajaran.
- Gunakan bahasa yang mudah dipahami peserta didik.
""",

    "11. Lembar Kerja Peserta Didik (LKPD)": """
FORMAT WAJIB LKPD:

1. Judul
2. Identitas peserta didik
3. Tujuan pembelajaran
4. Petunjuk belajar
5. Alat dan bahan
6. Tugas/langkah kerja
7. Pertanyaan
8. Lembar jawaban
9. Kesimpulan
10. Refleksi

ATURAN:
- Tugas harus dapat dikerjakan peserta didik.
- Instruksi harus jelas.
- Sesuaikan tingkat kesulitan dengan fase/kelas.
""",

    "12. Media Pembelajaran": """
FORMAT WAJIB:

Tabel:
| No | Media | Jenis | Tujuan Penggunaan | Cara Penggunaan | Materi | Sumber |

Dapat mencakup:
- PPT
- Video pembelajaran
- Kartu kelompok
- Infografis
- Alat peraga
- Media digital
- Media fisik

ATURAN:
- Media harus relevan dengan materi.
- Jangan mencantumkan tautan palsu.
""",

    "13. Asesmen Diagnostik": """
FORMAT WAJIB:

A. Diagnostik Non-Kognitif
- kondisi belajar
- minat
- motivasi
- kesiapan belajar
- kebutuhan belajar

B. Diagnostik Kognitif
- prasyarat materi
- pertanyaan/soal
- kunci jawaban
- pedoman interpretasi hasil

ATURAN:
- Instrumen harus siap digunakan.
- Sertakan petunjuk dan cara membaca hasil.
""",

    "14. Asesmen Formatif": """
FORMAT WAJIB:

1. Tujuan asesmen
2. Teknik asesmen
3. Instrumen
4. Lembar observasi/catatan
5. Soal atau tugas
6. Rubrik/pedoman penilaian
7. Kriteria ketercapaian
8. Tindak lanjut

Boleh menggunakan:
- observasi
- catatan anekdot
- penilaian diri
- penilaian teman
- kuis
- tugas
""",

    "15. Asesmen Sumatif": """
FORMAT WAJIB:

1. Identitas asesmen
2. Kisi-kisi soal
3. Indikator
4. Soal pilihan ganda dan/atau esai
5. Kunci jawaban
6. Pedoman penskoran
7. Rubrik
8. Analisis ketercapaian

ATURAN:
- Soal harus sesuai indikator.
- Tingkat kesulitan harus sesuai fase/kelas.
- Jangan membuat soal yang jawabannya ambigu.
""",

    "16. Daftar Nilai": """
FORMAT WAJIB:

A. Nilai Formatif per TP
B. Nilai Sumatif per Lingkup Materi
C. Nilai Akhir Semester
D. Rekapitulasi

Tabel minimal:
| No | Nama Siswa | TP 1 | TP 2 | TP 3 | Sumatif | Nilai Akhir | Keterangan |

ATURAN:
- Buat format yang siap diisi.
- Jangan mengarang nama atau nilai siswa.
""",

    "17. Jurnal Agenda Guru": """
FORMAT WAJIB:

| Hari/Tanggal | Kelas | Jam Ke- | Materi Pokok | Kehadiran Siswa | Catatan Kejadian |

Sertakan ruang untuk:
- refleksi singkat
- tindak lanjut
- catatan khusus
""",

    "18. Program Remedial & Pengayaan": """
FORMAT WAJIB:

Tabel:
| No | Nama Siswa | Nilai Awal | Tujuan/Kompetensi | Bentuk Intervensi | Nilai Akhir | Status |

Bentuk intervensi dapat berupa:
- tugas
- tutor sebaya
- pembelajaran ulang
- latihan tambahan
- pendampingan

Pisahkan program:
A. Remedial
B. Pengayaan
""",

    "19. Refleksi Pembelajaran": """
FORMAT WAJIB:

A. Apa yang berhasil?
B. Apa yang belum berhasil?
C. Kendala pembelajaran
D. Respons peserta didik
E. Bukti ketercapaian tujuan
F. Faktor penyebab
G. Rencana tindak lanjut
H. Perbaikan pembelajaran berikutnya
""",

    "20. Glosarium & Daftar Pustaka": """
FORMAT WAJIB:

A. GLOSARIUM

| No | Istilah | Pengertian |

B. DAFTAR PUSTAKA

Gunakan sumber yang benar-benar relevan.

ATURAN:
- Jangan membuat buku, penulis, penerbit, DOI,
  atau URL palsu.
- Jika sumber tidak tersedia, gunakan sumber umum
  yang dapat diverifikasi dan jangan mengarang detail.
""",

    "21. Buku Absensi Siswa": """
FORMAT WAJIB:

Tabel rekap kehadiran:
| No | Nama Siswa | Sakit | Izin | Alfa | Jumlah Kehadiran | Keterangan |

Sertakan format:
- daftar siswa
- rekap harian
- rekap semester

ATURAN:
- Jangan mengarang nama siswa.
- Sediakan baris kosong yang siap diisi.
""",

    "22. Dokumen Analisis Hasil Belajar": """
FORMAT WAJIB:

1. Identitas pembelajaran
2. Rekap hasil belajar
3. Analisis ketercapaian TP
4. Analisis per indikator
5. Materi yang belum tuntas
6. Peserta didik yang memerlukan tindak lanjut
7. Analisis ketuntasan
8. Remedial
9. Pengayaan
10. Kesimpulan
11. Rencana tindak lanjut

Tabel utama:
| No | Indikator/TP | Jumlah Siswa | Tuntas | Belum Tuntas | Persentase | Tindak Lanjut |

ATURAN:
- Jangan mengarang data siswa atau nilai.
- Gunakan format siap diisi jika data belum diberikan.
"""
}


# =========================================================
# PROMPT GEMINI
# =========================================================

def buat_prompt(data):

    profil = ", ".join(
        data["profil_pancasila"]
    )

    if not profil:
        profil = "Sesuai karakteristik materi"

    aturan = ATURAN_22.get(
        data["jenis_perangkat"],
        "Gunakan format administrasi pembelajaran yang sesuai."
    )

    return f"""
Anda adalah ahli penyusunan perangkat pembelajaran
Kurikulum Merdeka di Indonesia.

TUGAS UTAMA
Buat SATU dokumen perangkat pembelajaran sesuai pilihan
perangkat yang diberikan pengguna.

========================================================
JENIS PERANGKAT
========================================================

{data["jenis_perangkat"]}

========================================================
ATURAN KHUSUS PERANGKAT
========================================================

{aturan}

========================================================
IDENTITAS SATUAN PENDIDIKAN
========================================================

Dinas Pendidikan:
{data["dinas"]}

Satuan Pendidikan:
{data["sekolah"]}

Alamat:
{data["alamat"]}

Kota/Kabupaten:
{data["kota"]}

========================================================
IDENTITAS GURU
========================================================

Jabatan:
{data["jabatan_guru"]}

Nama Guru:
{data["guru_nama"]}

NIP Guru:
{data["guru_nip"]}

Kepala Sekolah:
{data["ks_nama"]}

NIP Kepala Sekolah:
{data["ks_nip"]}

Pengawas Pembina:
{data["pengawas_nama"]}

NIP Pengawas:
{data["pengawas_nip"]}

========================================================
PARAMETER PEMBELAJARAN
========================================================

Mata Pelajaran:
{data["mapel"]}

Fase/Kelas:
{data["fase_kelas"]}

Semester:
{data["semester"]}

Alokasi Waktu:
{data["alokasi_waktu"]}

Materi Pokok:
{data["materi_pokok"]}

Profil Pelajar Pancasila:
{profil}

Tanggal:
{data["tanggal"]}

========================================================
ATURAN UMUM
========================================================

1. Gunakan bahasa Indonesia formal, baku dan jelas.

2. Buat dokumen lengkap dan siap digunakan guru.

3. Ikuti ATURAN KHUSUS PERANGKAT di atas.

4. Jangan mencampurkan format perangkat lain.

5. Jangan mengarang data siswa, nilai, tanggal,
   nomor surat, nomor regulasi, NIP, nama pejabat,
   alamat, atau data resmi lainnya.

6. Jika data yang diperlukan belum tersedia, gunakan
   tabel/format yang siap diisi, bukan data palsu.

7. Jangan membuat klaim bahwa suatu dokumen adalah
   dokumen resmi pemerintah jika sumber resminya
   tidak diberikan.

8. Untuk CP atau regulasi pemerintah:
   jangan mengarang kutipan resmi.

9. Untuk kalender pendidikan:
   jangan mengarang hari libur atau pekan efektif resmi.

10. Untuk daftar pustaka:
    jangan membuat sumber palsu.

11. Gunakan heading Markdown:
    # untuk judul utama
    ## untuk bagian
    ### untuk subbagian

12. Gunakan tabel Markdown untuk data tabular.

13. Pastikan setiap tabel mempunyai header yang jelas.

14. Jangan menggunakan blok kode Markdown.

15. Jangan memberikan penjelasan di luar dokumen.

16. Jangan menggunakan placeholder seperti:
    "[isi]"
    "[sesuaikan]"
    "[masukkan]"
    "..."
    "dst."
    "dan seterusnya"

17. Jika dokumen membutuhkan pengesahan, buat bagian:

    PENGESAHAN

    Mengetahui,
    Kepala Sekolah

    Nama:
    NIP:

    Guru/Pendidik

    Nama:
    NIP:

    Pengawas Pembina
    jika relevan.

18. Jangan mengisi nama, NIP atau tanda tangan dengan
    data yang tidak diberikan.

19. Pastikan dokumen sesuai dengan:
    - fase
    - kelas
    - mata pelajaran
    - semester
    - materi pokok
    - alokasi waktu.

20. Periksa kembali konsistensi antara TP, materi,
    asesmen dan alokasi waktu sebelum menghasilkan
    dokumen.

========================================================
HASIL AKHIR
========================================================

Mulai langsung dari judul dokumen.

Hasil harus berupa dokumen yang dapat langsung
dipreview dan diunduh sebagai DOCX.

Jangan memberikan komentar kepada pengguna.
Jangan menjelaskan proses pembuatan.
Hanya hasil dokumen.
"""
    
# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:

    st.markdown(
        f"""
        <div class="brand-logo">
            <img
                src="{LOGO_URL}"
                alt="Logo SIAP AJAR 22"
            >
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown(
        "<h3 style='text-align:center;'>SIAP AJAR 22</h3>",
        unsafe_allow_html=True
    )

    st.caption("Portal Administrasi Pembelajaran")

    st.divider()

    st.markdown("### 👤 Pendidik Aktif")

    st.success(
        f"**{st.session_state.user_name}**\n\n"
        "🟢 Sesi AI aktif"
    )

    if st.button(
        "🚪 Keluar / Ganti Akun",
        use_container_width=True
    ):
        st.session_state.authenticated = False
        st.session_state.user_name = ""
        st.session_state.user_api_key = ""
        st.session_state.hasil_teks = ""

        st.rerun()

    st.divider()

    st.markdown("### ⚙️ Status Sistem")

    st.markdown(
        """
        **Engine:** Google Gemini

        **Output:** DOCX / TXT

        **Perangkat:** 22 Dokumen
        """
    )


# =========================================================
# HEADER
# =========================================================

st.title(
    f"Selamat Berkarya, {st.session_state.user_name}"
)

st.caption(
    "Susun perangkat pembelajaran secara otomatis, "
    "lengkap, sistematis, dan siap digunakan."
)


# =========================================================
# STATISTIK
# =========================================================

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown(
        """
        <div class="info-card">
            <div class="info-number">22</div>
            <div class="info-label">
                Perangkat Pembelajaran
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

with col2:
    st.markdown(
        """
        <div class="info-card">
            <div class="info-number">AI</div>
            <div class="info-label">
                Google Gemini
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

with col3:
    st.markdown(
        """
        <div class="info-card">
            <div class="info-number">DOCX</div>
            <div class="info-label">
                Format Siap Unduh
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )


# =========================================================
# KONFIGURASI PEMBELAJARAN
# =========================================================

with st.container(border=True):

    st.markdown(
        '<div class="section-title">'
        '📋 1. Konfigurasi Pembelajaran'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2 = st.columns(2)

    with col1:

        fase_kelas = st.selectbox(
            "Tingkatan Kelas / Fase",
            [
                "Fase A - Kelas 1 SD/MI",
                "Fase A - Kelas 2 SD/MI",
                "Fase B - Kelas 3 SD/MI",
                "Fase B - Kelas 4 SD/MI",
                "Fase C - Kelas 5 SD/MI",
                "Fase C - Kelas 6 SD/MI",
                "Fase D - Kelas 7 SMP/MTs",
                "Fase D - Kelas 8 SMP/MTs",
                "Fase D - Kelas 9 SMP/MTs",
                "Fase E - Kelas 10 SMA/MA/SMK",
                "Fase F - Kelas 11 SMA/MA/SMK",
                "Fase F - Kelas 12 SMA/MA/SMK",
            ],
            index=3
        )

        is_sd = "SD/MI" in fase_kelas

        hasil_kelas = re.search(
            r"Kelas (\d+)",
            fase_kelas
        )

        angka_kelas = (
            hasil_kelas.group(1)
            if hasil_kelas
            else ""
        )

        if is_sd:
            jabatan_guru = (
                f"Guru Kelas {angka_kelas}"
                if angka_kelas
                else "Guru Kelas"
            )

            label_mapel = (
                "Mata Pelajaran / Muatan Pelajaran"
            )

            default_mapel = "IPAS"

        else:
            jabatan_guru = "Guru Mata Pelajaran"
            label_mapel = "Mata Pelajaran"
            default_mapel = "Fisika"

        mapel = st.text_input(
            label_mapel,
            value=default_mapel
        )

        materi_pokok = st.text_input(
            "Fokus Topik / Materi Pokok",
            placeholder="Contoh: Usaha dan Energi"
        )

    with col2:

        alokasi_waktu = st.text_input(
            "Alokasi Waktu / Target JP",
            value=(
                "4 JP / Minggu"
                if is_sd
                else "3 JP / Minggu"
            )
        )

        profil_pancasila = st.multiselect(
            "Dimensi Profil Pelajar Pancasila",
            [
                "Beriman, Bertakwa kepada Tuhan YME & "
                "Berakhlak Mulia",
                "Berkebinekaan Global",
                "Gotong Royong",
                "Mandiri",
                "Bernalar Kritis",
                "Kreatif",
            ],
            default=[
                "Bernalar Kritis",
                "Gotong Royong",
                "Mandiri",
            ]
        )


# =========================================================
# 22 PERANGKAT
# =========================================================

daftar_22_perangkat = [
    "01. Capaian Pembelajaran (CP)",
    "02. Alur Tujuan Pembelajaran (ATP)",
    "03. Modul Ajar (RPP Plus)",
    "04. Kriteria Ketercapaian Tujuan Pembelajaran (KKTP)",
    "05. Modul Projek (P5)",
    "06. Kalender Pendidikan",
    "07. Program Tahunan (Prota)",
    "08. Program Semester (Prosem)",
    "09. Jadwal Mengajar",
    "10. Bahan Ajar",
    "11. Lembar Kerja Peserta Didik (LKPD)",
    "12. Media Pembelajaran",
    "13. Asesmen Diagnostik",
    "14. Asesmen Formatif",
    "15. Asesmen Sumatif",
    "16. Daftar Nilai",
    "17. Jurnal Agenda Guru",
    "18. Program Remedial & Pengayaan",
    "19. Refleksi Pembelajaran",
    "20. Glosarium & Daftar Pustaka",
    "21. Buku Absensi Siswa",
    "22. Dokumen Analisis Hasil Belajar",
]
jenis_perangkat = st.selectbox(
    "📄 Pilih Dokumen yang Akan Dibuat",
    daftar_22_perangkat
)


# =========================================================
# SEMESTER
# =========================================================

semester = st.radio(
    "Semester Berjalan",
    [
        "Semester Ganjil",
        "Semester Genap",
        "Semester Ganjil & Genap (1 Tahun Penuh)",
    ],
    horizontal=True
)


# =========================================================
# IDENTITAS
# =========================================================

tab1, tab2 = st.tabs(
    [
        "🏫 Identitas Satuan Pendidikan",
        "✍️ Pejabat & Penandatangan",
    ]
)


with tab1:

    col1, col2 = st.columns(2)

    with col1:

        dinas = st.text_input(
            "Dinas Pendidikan Pembina",
            value=(
                "DINAS PENDIDIKAN KOTA PONTIANAK"
                if is_sd
                else
                "DINAS PENDIDIKAN PROVINSI "
                "KALIMANTAN BARAT"
            )
        )

        sekolah = st.text_input(
            "Nama Satuan Pendidikan",
            value=(
                "SDN 01 PONTIANAK"
                if is_sd
                else
                "SMAS NUSA HARAPAN"
            )
        )

    with col2:

        alamat = st.text_input(
            "Alamat & Kontak Sekolah",
            value=(
                "Jl. Pancasila No. 10, "
                "Telp. (0561) 734567"
            )
        )

        kota = st.text_input(
            "Kota / Kabupaten Domisili",
            value="Pontianak"
        )


# =========================================================
# PENANDATANGAN
# =========================================================

with tab2:

    col1, col2, col3 = st.columns(3)

    with col1:

        st.markdown(
            f"**{jabatan_guru}**"
        )

        guru_nama = st.text_input(
            "Nama Guru",
            value=st.session_state.user_name,
            key="guru_nama"
        )

        guru_nip = st.text_input(
            "NIP Guru",
            value="-",
            key="guru_nip"
        )

    with col2:

        st.markdown("**Kepala Sekolah**")

        ks_nama = st.text_input(
            "Nama Kepala Sekolah",
            placeholder="Contoh: Zulkifli, S.Pd.",
            key="ks_nama"
        )

        ks_nip = st.text_input(
            "NIP Kepala Sekolah",
            placeholder="Contoh: 197508122005011004",
            key="ks_nip"
        )

    with col3:

        st.markdown("**Pengawas Pembina**")

        pengawas_nama = st.text_input(
            "Nama Pengawas Pembina",
            placeholder="Contoh: Andar",
            key="pengawas_nama"
        )

        pengawas_nip = st.text_input(
            "NIP Pengawas",
            placeholder="Contoh: 196811231993032003",
            key="pengawas_nip"
        )


# =========================================================
# TERBITKAN
# =========================================================

st.divider()

terbitkan = st.button(
    "✨ TERBITKAN DOKUMEN ADMINISTRASI RESMI",
    use_container_width=True,
    type="primary"
)


if terbitkan:

    if not mapel.strip():
        st.warning("⚠️ Mata pelajaran wajib diisi.")
        st.stop()

    if not sekolah.strip():
        st.warning("⚠️ Nama sekolah wajib diisi.")
        st.stop()

    if not materi_pokok.strip():
        st.warning(
            "⚠️ Fokus topik/materi pokok wajib diisi."
        )
        st.stop()

    if not st.session_state.user_api_key.strip():
        st.error(
            "❌ API Key Gemini belum tersedia. "
            "Silakan kembali ke halaman login."
        )
        st.stop()

    data = {
        "dinas": dinas,
        "sekolah": sekolah,
        "alamat": alamat,
        "kota": kota,
        "jabatan_guru": jabatan_guru,
        "guru_nama": guru_nama,
        "guru_nip": guru_nip,
        "ks_nama": ks_nama or "-",
        "ks_nip": ks_nip or "-",
        "pengawas_nama": pengawas_nama or "-",
        "pengawas_nip": pengawas_nip or "-",
        "mapel": mapel,
        "fase_kelas": fase_kelas,
        "semester": semester,
        "alokasi_waktu": alokasi_waktu,
        "materi_pokok": materi_pokok,
        "profil_pancasila": profil_pancasila,
        "jenis_perangkat": jenis_perangkat,
        "tanggal": datetime.now().strftime("%d-%m-%Y"),
    }

    prompt = buat_prompt(data)

    progress = st.progress(0)
    status = st.empty()

    try:

        status.info(
            "⚡ Menghubungkan ke Google Gemini..."
        )

        client = genai.Client(
            api_key=st.session_state.user_api_key
        )

        progress.progress(15)

        model_list = [
            "gemini-3.6-flash"
        ]

        response = None
        model_berhasil = ""
        errors = []

        for model_name in model_list:

            for percobaan in range(3):

                try:

                    status.info(
                        f"📝 Membuat dokumen dengan "
                        f"{model_name} "
                        f"({percobaan + 1}/3)"
                    )

                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt
                    )

                    if (
                        response
                        and getattr(response, "text", None)
                    ):
                        model_berhasil = model_name
                        break

                except Exception as exc:

                    error_text = str(exc)

                    errors.append(
                        f"{model_name}: {error_text}"
                    )

                    if (
                        "503" in error_text
                        or "UNAVAILABLE" in error_text
                    ):

                        tunggu = 2 ** percobaan

                        status.warning(
                            f"⏳ Gemini sedang sibuk. "
                            f"Mencoba lagi dalam "
                            f"{tunggu} detik..."
                        )

                        time.sleep(tunggu)

                    else:
                        break

            if (
                response
                and getattr(response, "text", None)
            ):
                break

        progress.progress(90)

        if (
            not response
            or not getattr(response, "text", None)
        ):
            raise RuntimeError(
                "Semua percobaan Gemini gagal.\n\n"
                + "\n".join(errors[-6:])
            )

        st.session_state.hasil_teks = response.text

        nama_bersih = re.sub(
            r"^\d+\.\s*",
            "",
            jenis_perangkat
        )

        st.session_state.nama_file_base = re.sub(
            r"[^A-Za-z0-9_-]+",
            "_",
            nama_bersih
        ).strip("_")

        progress.progress(100)

        status.success(
            f"✅ Dokumen berhasil dibuat dengan "
            f"{model_berhasil}."
        )

        time.sleep(0.5)

        progress.empty()
        status.empty()

    except Exception as exc:

        progress.empty()
        status.empty()

        st.error(
            "❌ Gagal menerbitkan dokumen."
        )

        st.code(
            str(exc),
            language="text"
        )


# =========================================================
# PREVIEW
# =========================================================

if st.session_state.hasil_teks:

    st.divider()

    st.markdown(
        "### 📄 Preview Lembar Kerja A4"
    )

    preview = markdown_to_html(
        st.session_state.hasil_teks
    )

    st.markdown(
        f"""
        <div class="paper-a4">
            {preview}
        </div>
        """,
        unsafe_allow_html=True
    )

    col1, col2 = st.columns(2)

    with col1:

        docx_file = buat_file_docx(
            st.session_state.hasil_teks
        )

        st.download_button(
            "📄 Unduh Berkas Word (.DOCX)",
            data=docx_file,
            file_name=(
                f"{st.session_state.nama_file_base}.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True
        )

    with col2:

        st.download_button(
            "📝 Unduh Teks Mentah (.TXT)",
            data=st.session_state.hasil_teks,
            file_name=(
                f"{st.session_state.nama_file_base}.txt"
            ),
            mime="text/plain",
            use_container_width=True
        )


# =========================================================
# FOOTER
# =========================================================

st.markdown(
    """
    <div class="footer-box">
        SIAP AJAR 22 • Creator: Andar<br>
        © 2026 SIAP AJAR 22 • Engine AI Pembelajaran
    </div>
    """,
    unsafe_allow_html=True
)
