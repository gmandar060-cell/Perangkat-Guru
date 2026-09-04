import streamlit as st
from google import genai
from datetime import datetime
import re
import io
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# =========================================================
# SIAP AJAR 22 - APP UTAMA
# =========================================================

st.set_page_config(
    page_title="SIAP AJAR 22 | Portal Administrasi Kurikulum",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)


# =========================================================
# CSS
# =========================================================

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"], .stMarkdown, p, span, label,
h1, h2, h3, h4, h5, h6, button, input, textarea {
    font-family: 'Plus Jakarta Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
}

.stApp {
    background: #f8fafc !important;
}

.block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 3rem !important;
    max-width: 1120px !important;
}


/* =========================================================
   LOGIN
   ========================================================= */

.stApp:has(.login-page) {
    background:
        linear-gradient(rgba(255,255,255,.055) 1px, transparent 1px),
        linear-gradient(90deg, rgba(255,255,255,.055) 1px, transparent 1px),
        linear-gradient(135deg, #2547a8 0%, #1f45ad 48%, #2d5bd0 100%) !important;
    background-size: 28px 28px, 28px 28px, 100% 100% !important;
}

.stApp:has(.login-page) .block-container {
    max-width: 900px !important;
    padding-top: 26px !important;
    padding-bottom: 24px !important;
}

.login-page {
    color: white;
}

.login-hero {
    text-align: center;
    color: white;
    margin: 0 auto 28px auto;
}

.login-icon {
    width: 80px;
    height: 80px;
    margin: 0 auto 14px auto;
    border-radius: 24px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 38px;
    background: rgba(255,255,255,.10);
    border: 1px solid rgba(255,255,255,.30);
    box-shadow: 0 10px 30px rgba(0,0,0,.12);
}

.login-badge {
    display: inline-block;
    padding: 8px 16px;
    margin-left: 8px;
    border: 1px solid rgba(255,255,255,.32);
    border-radius: 999px;
    color: rgba(255,255,255,.90);
    font-size: 12px;
    font-weight: 700;
    letter-spacing: .2px;
    vertical-align: middle;
    background: rgba(255,255,255,.08);
}

.login-title {
    margin: 0;
    color: #fff !important;
    font-size: clamp(38px, 6vw, 52px) !important;
    line-height: 1.05;
    font-weight: 800;
    letter-spacing: -2px;
}

.login-title span {
    color: #74a9ff !important;
}

.login-subtitle {
    margin: 10px auto 0;
    max-width: 520px;
    color: rgba(255,255,255,.88) !important;
    font-size: 16px;
    line-height: 1.5;
}

.login-subtitle strong {
    color: #fff !important;
}

.stat-card {
    min-height: 88px;
    border-radius: 16px;
    padding: 14px 10px;
    box-sizing: border-box;
    text-align: center;
    color: white;
    background: rgba(219,228,255,.82);
    border: 1px solid rgba(255,255,255,.28);
    box-shadow: 0 8px 22px rgba(11,36,108,.12);
    backdrop-filter: blur(8px);
}

.stat-icon {
    font-size: 21px;
    line-height: 1;
    margin-bottom: 7px;
}

.stat-number {
    font-size: 14px;
    font-weight: 800;
    color: #fff !important;
    margin: 0;
}

.stat-label {
    font-size: 11px;
    color: rgba(255,255,255,.72) !important;
    margin-top: 2px;
}

.login-form-title {
    color: #172033 !important;
    font-size: 19px;
    font-weight: 800;
    margin: 0 0 2px 0;
}

.login-form-desc {
    color: #94a3b8 !important;
    font-size: 13px;
    margin: 0 0 18px 0;
}

.stApp:has(.login-page) [data-testid="stForm"] {
    background: #fff !important;
    border: 1px solid rgba(255,255,255,.75) !important;
    border-radius: 18px !important;
    padding: 28px 30px !important;
    margin: 24px auto 0 !important;
    box-shadow: 0 18px 45px rgba(11,31,88,.20) !important;
    box-sizing: border-box !important;
    max-width: 100% !important;
}

.stApp:has(.login-page) [data-testid="stForm"] .stTextInput label {
    color: #334155 !important;
    font-size: 13px !important;
    font-weight: 700 !important;
}

.stApp:has(.login-page) [data-testid="stForm"] .stTextInput input {
    background: #fff !important;
    color: #1e293b !important;
    border: 1px solid #dbe3ef !important;
    border-radius: 11px !important;
    min-height: 48px !important;
    font-size: 14px !important;
}

.stApp:has(.login-page) [data-testid="stForm"] .stTextInput input::placeholder {
    color: #c2ccda !important;
    -webkit-text-fill-color: #c2ccda !important;
}

.stApp:has(.login-page) [data-testid="stForm"] .stButton > button {
    margin-top: 10px !important;
    min-height: 52px !important;
    border: none !important;
    border-radius: 11px !important;
    background: linear-gradient(90deg,#2760df 0%,#2d68ee 100%) !important;
    box-shadow: 0 8px 18px rgba(37,99,235,.24) !important;
}

.stApp:has(.login-page) [data-testid="stForm"] .stButton > button,
.stApp:has(.login-page) [data-testid="stForm"] .stButton > button * {
    color: white !important;
    font-weight: 800 !important;
    font-size: 15px !important;
}

.api-help {
    margin: 7px 0 2px;
    font-size: 12px;
}

.api-help a {
    color: #2563eb !important;
    text-decoration: none !important;
    font-weight: 700;
}

.login-footer {
    text-align: center;
    color: rgba(255,255,255,.80);
    font-size: 12px;
    margin: 22px 0 0;
}

.creator-pill {
    display: inline-block;
    padding: 7px 15px;
    border: 1px solid rgba(255,255,255,.18);
    border-radius: 999px;
    background: rgba(255,255,255,.06);
}

.creator-dot {
    display: inline-block;
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background: #42d3b2;
    margin-right: 6px;
}


/* =========================================================
   DASHBOARD
   ========================================================= */

.stAlert {
    border-radius: 10px !important;
}

.app-header {
    background: linear-gradient(135deg, #162b69 0%, #2347a7 58%, #2d68ee 100%) !important;
    border-radius: 18px;
    padding: 28px 32px;
    margin-bottom: 24px;
    box-shadow: 0 12px 30px rgba(15, 23, 42, .12);
}

.app-header h1 {
    font-size: 26px;
    font-weight: 800;
    margin: 0 0 6px 0;
    color: #fff !important;
}

.app-header p {
    font-size: 14px;
    color: #dbeafe !important;
    margin: 0;
    line-height: 1.5;
}

details[data-testid="stExpander"] {
    border: 1px solid #e2e8f0 !important;
    border-radius: 10px !important;
    background: #fff !important;
}

input, textarea, select {
    color: #0f172a !important;
    background-color: #fff !important;
    -webkit-text-fill-color: #0f172a !important;
}

input::placeholder,
textarea::placeholder {
    color: #94a3b8 !important;
    -webkit-text-fill-color: #94a3b8 !important;
    opacity: 1 !important;
}

.stTextInput label,
.stSelectbox label,
.stMultiSelect label,
.stRadio label {
    color: #1e293b !important;
    font-weight: 600 !important;
}

.stButton > button {
    background: linear-gradient(135deg, #1e3a8a 0%, #2563eb 100%) !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 12px 24px !important;
    box-shadow: 0 4px 12px rgba(37,99,235,.25) !important;
    width: 100% !important;
}

.stButton > button,
.stButton > button * {
    color: #fff !important;
    font-weight: 700 !important;
    font-size: 15px !important;
}

.stDownloadButton > button {
    font-weight: 700 !important;
    border-radius: 10px !important;
    padding: 12px 20px !important;
    width: 100% !important;
}

.paper-a4 {
    background: #fff !important;
    border: 1px solid #cbd5e1;
    box-shadow: 0 12px 30px rgba(0,0,0,.08);
    padding: 44px 50px;
    margin: 16px auto;
    border-radius: 4px;
    max-width: 900px;
    color: #0f172a !important;
}

.paper-a4 table {
    width: 100% !important;
    border-collapse: collapse !important;
    margin: 14px 0 !important;
    font-size: 13px !important;
}

.paper-a4 th {
    background: #f1f5f9 !important;
    border: 1px solid #475569 !important;
    padding: 8px 10px !important;
    text-align: left;
    font-weight: 700;
}

.paper-a4 td {
    border: 1px solid #64748b !important;
    padding: 7px 10px !important;
}

.footer-box {
    text-align: center;
    padding: 24px 10px 10px;
    color: #64748b !important;
    font-size: 12px;
    border-top: 1px solid #e2e8f0;
    margin-top: 40px;
}

section[data-testid="stSidebar"] {
    background-color: #fff !important;
    border-right: 1px solid #e2e8f0;
}

@media (max-width: 768px) {
    .stApp:has(.login-page) .block-container {
        padding: 18px 14px 24px !important;
    }

    .login-icon {
        width: 68px;
        height: 68px;
        font-size: 31px;
    }

    .login-title {
        font-size: 38px !important;
    }

    .login-subtitle {
        font-size: 14px;
    }

    .paper-a4 {
        padding: 20px 16px;
    }

    .app-header {
        padding: 22px 18px;
    }

    .stat-card {
        min-height: 78px;
        padding: 11px 5px;
    }
}
</style>
""",
    unsafe_allow_html=True,
)


# =========================================================
# LOGO
# =========================================================

LOGO_URL = (
    "https://cdn.jsdelivr.net/gh/twitter/twemoji@14.0.2/"
    "assets/72x72/1f393.png"
)


def render_logo(width=90):
    st.markdown(
        f"""
        <div style="
            display:flex;
            justify-content:center;
            align-items:center;
            margin-bottom:12px;
        ">
            <img
                src="{LOGO_URL}"
                width="{width}"
                height="{width}"
                alt="Logo SIAP AJAR 22"
                style="filter:drop-shadow(0 4px 6px rgba(0,0,0,.1));"
            >
        </div>
        """,
        unsafe_allow_html=True,
    )


# =========================================================
# SESSION STATE
# =========================================================

defaults = {
    "authenticated": False,
    "user_name": "",
    "user_api_key": "",
    "hasil_teks": "",
    "nama_file_base": "Perangkat_Kurikulum",
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# =========================================================
# EXPORT WORD
# =========================================================

def buat_file_docx(markdown_text: str) -> io.BytesIO:

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):

        line = lines[i].strip()

        if line.startswith("```"):
            i += 1
            continue

        # TABLE
        if line.startswith("|") and line.endswith("|"):

            table_lines = []

            while (
                i < len(lines)
                and lines[i].strip().startswith("|")
                and lines[i].strip().endswith("|")
            ):

                raw = lines[i].strip()

                if not re.match(r"^\|[\s\-:|]+\|$", raw):

                    cells = [
                        c.strip()
                        for c in raw[1:-1].split("|")
                    ]

                    table_lines.append(cells)

                i += 1

            if table_lines:

                rows = len(table_lines)
                cols = max(len(row) for row in table_lines)

                table = doc.add_table(
                    rows=rows,
                    cols=cols
                )

                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                for r_idx, row_data in enumerate(table_lines):

                    for c_idx in range(cols):

                        value = (
                            row_data[c_idx]
                            if c_idx < len(row_data)
                            else ""
                        )

                        value = (
                            value
                            .replace("<br>", "\n")
                            .replace("<br/>", "\n")
                            .replace("**", "")
                        )

                        cell = table.cell(
                            r_idx,
                            c_idx
                        )

                        cell.text = value

                        for paragraph in cell.paragraphs:

                            for run in paragraph.runs:

                                run.font.name = "Calibri"
                                run.font.size = Pt(9.5)

                                if r_idx == 0:
                                    run.font.bold = True

                        if r_idx == 0:

                            shading = parse_xml(
                                r'<w:shd {} w:fill="E2E8F0"/>'
                                .format(nsdecls("w"))
                            )

                            cell._tc.get_or_add_tcPr().append(
                                shading
                            )

                borders = parse_xml(
                    f"""
                    <w:tblBorders {nsdecls("w")}>
                        <w:top w:val="single" w:sz="4" w:color="94A3B8"/>
                        <w:bottom w:val="single" w:sz="4" w:color="94A3B8"/>
                        <w:insideH w:val="single" w:sz="4" w:color="CBD5E1"/>
                        <w:insideV w:val="single" w:sz="4" w:color="CBD5E1"/>
                    </w:tblBorders>
                    """
                )

                table._tbl.tblPr.append(borders)
                doc.add_paragraph()

            continue

        if line.startswith("# "):

            p = doc.add_heading(
                line[2:],
                level=1
            )

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith("## "):

            doc.add_heading(
                line[3:],
                level=2
            )

        elif line.startswith("### "):

            doc.add_heading(
                line[4:],
                level=3
            )

        elif line.startswith("---"):

            p = doc.add_paragraph()

            border = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" '
                f'w:space="1" w:color="000000"/>'
                f'</w:pBdr>'
            )

            p._p.get_or_add_pPr().append(border)

        elif line:

            clean = re.sub(
                r"<.*?>",
                "",
                line
            )

            clean = re.sub(
                r"\*\*(.*?)\*\*",
                r"\1",
                clean
            )

            p = doc.add_paragraph(clean)

            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)

            for run in p.runs:

                run.font.name = "Calibri"
                run.font.size = Pt(11)

        i += 1

    stream = io.BytesIO()

    doc.save(stream)

    stream.seek(0)

    return stream


# =========================================================
# MARKDOWN -> HTML
# =========================================================

def markdown_to_html(text: str) -> str:

    lines = text.splitlines()

    html = []

    in_table = False
    table_rows = []

    def flush_table():

        nonlocal table_rows, in_table

        if not table_rows:
            return

        html.append("<table>")

        for r_idx, row in enumerate(table_rows):

            tag = "th" if r_idx == 0 else "td"

            html.append("<tr>")

            for cell in row:

                cell_html = (
                    cell
                    .replace("<br/>", "<br>")
                    .replace("<br>", "<br>")
                )

                html.append(
                    f"<{tag}>{cell_html}</{tag}>"
                )

            html.append("</tr>")

        html.append("</table>")

        table_rows = []
        in_table = False

    for raw in lines:

        line = raw.strip()

        if line.startswith("```"):
            continue

        if line.startswith("|") and line.endswith("|"):

            if re.match(
                r"^\|[\s\-:|]+\|$",
                line
            ):
                continue

            in_table = True

            cells = [
                c.strip()
                for c in line[1:-1].split("|")
            ]

            table_rows.append(cells)

            continue

        if in_table:
            flush_table()

        if not line:

            html.append("<p>&nbsp;</p>")

            continue

        escaped = (
            line
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        escaped = re.sub(
            r"\*\*(.*?)\*\*",
            r"<strong>\1</strong>",
            escaped
        )

        if escaped.startswith("### "):

            html.append(
                f"<h3>{escaped[4:]}</h3>"
            )

        elif escaped.startswith("## "):

            html.append(
                f"<h2>{escaped[3:]}</h2>"
            )

        elif escaped.startswith("# "):

            html.append(
                f"<h1>{escaped[2:]}</h1>"
            )

        elif escaped.startswith("---"):

            html.append("<hr>")

        else:

            html.append(
                f"<p>{escaped}</p>"
            )

    if in_table:
        flush_table()

    return "\n".join(html)


# =========================================================
# PROMPT
# =========================================================

def buat_instruksi_prompt(data: dict) -> str:

    p3 = (
        ", ".join(data["profil_pancasila"])
        if data["profil_pancasila"]
        else "Sesuai karakteristik materi"
    )

    tanda_tangan = (
        "| Mengetahui,<br>Pengawas Pembina | "
        "Mengetahui,<br>Kepala Sekolah | "
        + data["kota_sekolah"]
        + ", "
        + data["tanggal_hari_ini"]
        + "<br>"
        + data["jabatan_guru"]
        + " |\n"
        "| :---: | :---: | :---: |\n"
        "| <br><br><br><br> | <br><br><br><br> | <br><br><br><br> |\n"
        "| **"
        + data["pengawas_nama"]
        + "**<br>NIP. "
        + data["pengawas_nip"]
        + " | **"
        + data["ks_nama"]
        + "**<br>NIP. "
        + data["ks_nip"]
        + " | **"
        + data["guru_nama"]
        + "**<br>NIP. "
        + data["guru_nip"]
        + " |"
    )

    return (
        "Bertindaklah sebagai ahli penyusunan perangkat "
        "pembelajaran Indonesia yang memahami Kurikulum "
        "Merdeka dan regulasi pendidikan yang berlaku.\n\n"

        "Tugas: susun dokumen \""
        + data["jenis_perangkat"]
        + "\" yang lengkap, detail, operasional, "
        "sistematis, dan siap disalin ke dokumen resmi sekolah.\n\n"

        "IDENTITAS SATUAN PENDIDIKAN\n"
        "- Dinas Pendidikan: "
        + data["dinas"]
        + "\n"
        "- Satuan Pendidikan: "
        + data["sekolah"]
        + "\n"
        "- Alamat dan Kontak: "
        + data["alamat"]
        + "\n"
        "- Kota/Kabupaten: "
        + data["kota_sekolah"]
        + "\n\n"

        "IDENTITAS PENDIDIK DAN PEJABAT\n"
        "- "
        + data["jabatan_guru"]
        + ": "
        + data["guru_nama"]
        + " (NIP: "
        + data["guru_nip"]
        + ")\n"
        "- Kepala Sekolah: "
        + data["ks_nama"]
        + " (NIP: "
        + data["ks_nip"]
        + ")\n"
        "- Pengawas Pembina: "
        + data["pengawas_nama"]
        + " (NIP: "
        + data["pengawas_nip"]
        + ")\n\n"

        "PARAMETER PEMBELAJARAN\n"
        "- Mata Pelajaran: "
        + data["mapel"]
        + "\n"
        "- Fase/Kelas: "
        + data["fase_kelas"]
        + "\n"
        "- Semester: "
        + data["semester"]
        + "\n"
        "- Alokasi Waktu: "
        + data["alokasi_waktu"]
        + "\n"
        "- Materi Pokok: "
        + data["materi_pokok"]
        + "\n"
        "- Dimensi Profil Pelajar Pancasila: "
        + p3
        + "\n"
        "- Tanggal: "
        + data["tanggal_hari_ini"]
        + "\n\n"

        "ATURAN OUTPUT\n"
        "1. Langsung mulai dengan KOP SURAT resmi lembaga pendidikan.\n"
        "2. Gunakan judul dokumen yang jelas menggunakan heading Markdown.\n"
        "3. Gunakan tabel Markdown yang rapi untuk identitas dokumen, rincian materi, atau kisi-kisi.\n"
        "4. DILARANG menggunakan simbol '...', '[lanjutkan]', '[sesuaikan]', 'dst.', atau 'dan seterusnya'.\n"
        "5. Jangan mengarang nomor regulasi atau undang-undang spesifik jika tidak yakin.\n"
        "6. Gunakan bahasa Indonesia formal, baku, edukatif, dan mudah diedit oleh guru.\n"
        "7. Jangan menyertakan blok kode Markdown berpagar.\n"
        "8. Akhiri dokumen dengan lembar pengesahan tiga kolom berikut:\n\n"
        + tanda_tangan
        + "\n\n"
        "Jangan menambahkan salam pembuka atau kalimat percakapan di luar dokumen."
    )


# =========================================================
# LOGIN
# =========================================================

if not st.session_state.authenticated:

    st.markdown(
        '<div class="login-page"></div>',
        unsafe_allow_html=True
    )

    # HERO
    st.markdown(
        """
        <div class="login-hero">
            <div style="
                display:flex;
                justify-content:center;
                align-items:center;
                gap:0;
            ">
                <div class="login-icon">🎓</div>
                <div class="login-badge">
                    ✦ &nbsp; KURIKULUM MERDEKA
                </div>
            </div>

            <h1 class="login-title">
                SIAP AJAR <span>22</span>
            </h1>

            <div class="login-subtitle">
                Satu Portal, Solusi Lengkap
                <strong>22 Perangkat Pembelajaran</strong>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # STATISTIK
    s1, s2, s3 = st.columns(
        3,
        gap="small"
    )

    with s1:
        st.markdown(
            """
            <div class="stat-card">
                <div class="stat-icon">▱</div>
                <div class="stat-number">4 Kategori</div>
                <div class="stat-label">Terorganisir</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with s2:
        st.markdown(
            """
            <div class="stat-card">
                <div class="stat-icon">▤</div>
                <div class="stat-number">22 Dokumen</div>
                <div class="stat-label">AI Generated</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with s3:
        st.markdown(
            """
            <div class="stat-card">
                <div class="stat-icon">▣</div>
                <div class="stat-number">Gemini AI</div>
                <div class="stat-label">Flash Model</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # FORM LOGIN
    with st.form(
        "login_form",
        clear_on_submit=False,
        border=False
    ):

        st.markdown(
            """
            <div class="login-form-title">
                Masuk ke Portal
            </div>

            <div class="login-form-desc">
                Isi data berikut untuk memulai
            </div>
            """,
            unsafe_allow_html=True,
        )

        nama_guru_input = st.text_input(
            "Nama Lengkap & Gelar",
            placeholder="Contoh: Andar Prasetyo, S.Pd.",
            key="login_nama",
        )

        api_key_masuk = st.text_input(
            "Gemini API Key",
            type="password",
            placeholder="AIza...",
            key="login_api_key",
        )

        st.markdown(
            """
            <div class="api-help">
                <a href="https://aistudio.google.com/apikey"
                   target="_blank">
                    ⓘ &nbsp; Bagaimana cara mendapatkan API key gratis?
                </a>
            </div>
            """,
            unsafe_allow_html=True,
        )

        masuk = st.form_submit_button(
            "Masuk ke Portal  →",
            use_container_width=True
        )

    if masuk:

        if not nama_guru_input.strip():

            st.warning(
                "⚠️ Mohon isi Nama Lengkap & Gelar terlebih dahulu."
            )

        elif not api_key_masuk.strip():

            st.warning(
                "⚠️ Mohon isi Gemini API Key terlebih dahulu."
            )

        else:

            st.session_state.user_name = (
                nama_guru_input.strip()
            )

            st.session_state.user_api_key = (
                api_key_masuk.strip()
            )

            st.session_state.authenticated = True

            st.rerun()

    st.markdown(
        """
        <div class="login-footer">
            <span class="creator-pill">
                <span class="creator-dot"></span>
                Creator: <strong>Andar</strong>
            </span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.stop()


# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:

    render_logo(70)

    st.markdown(
        """
        <h3 style="
            text-align:center;
            margin:2px 0;
            font-size:16px;
        ">
            SIAP AJAR 22
        </h3>
        """,
        unsafe_allow_html=True,
    )

    st.caption(
        "Portal Administrasi Pembelajaran"
    )

    st.divider()

    st.markdown(
        "#### 👤 Pendidik Aktif"
    )

    st.success(
        f"**{st.session_state.user_name}**\n\n"
        "🟢 Sesi AI aktif"
    )

    if st.button(
        "🔄 Keluar / Ganti Akun",
        use_container_width=True
    ):

        st.session_state.authenticated = False
        st.session_state.user_name = ""
        st.session_state.user_api_key = ""
        st.session_state.hasil_teks = ""

        st.rerun()

    st.divider()

    st.markdown(
        """
        **⚙️ Status Sistem**

        - **Engine:** Google Gemini
        - **Output:** DOCX / TXT
        - **Perangkat:** 22 dokumen
        """
    )


# =========================================================
# DASHBOARD
# =========================================================

st.markdown(
    f"""
    <div class="app-header">
        <h1>
            Selamat Berkarya,
            {st.session_state.user_name}
        </h1>

        <p>
            Susun perangkat pembelajaran secara terstruktur
            dan ekspor langsung ke Microsoft Word melalui
            platform SIAP AJAR 22.
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# KONFIGURASI
# =========================================================

with st.container(border=True):

    st.markdown(
        "### 📋 1. Konfigurasi Kurikulum & Jenjang"
    )

    col1, col2 = st.columns(2)

    with col1:

        fase_kelas = st.selectbox(
            "Tingkatan Kelas / Fase",
            [
                "Fase A - Kelas 1 SD/MI",
                "Fase A - Kelas 2 SD/MI",
                "Fase B - Kelas 3 SD/MI",
                "Fase B - Kelas 4 SD/MI",
                "Fase C - Kelas 5 SD/MI",
                "Fase C - Kelas 6 SD/MI",
                "Fase D - Kelas 7 SMP/MTs",
                "Fase D - Kelas 8 SMP/MTs",
                "Fase D - Kelas 9 SMP/MTs",
                "Fase E - Kelas 10 SMA/MA/SMK",
                "Fase F - Kelas 11 SMA/MA/SMK",
                "Fase F - Kelas 12 SMA/MA/SMK",
            ],
            index=3,
        )

        is_sd = "SD/MI" in fase_kelas

        if is_sd:

            match = re.search(
                r"Kelas (\d+)",
                fase_kelas
            )

            angka_kelas = (
                match.group(1)
                if match
                else ""
            )

            jabatan_guru_otomatis = (
                f"Guru Kelas {angka_kelas}"
                if angka_kelas
                else "Guru Kelas"
            )

            label_mapel = (
                "Mata Pelajaran / Muatan Pelajaran"
            )

            default_mapel = "IPAS"

        else:

            jabatan_guru_otomatis = (
                "Guru Mata Pelajaran"
            )

            label_mapel = "Mata Pelajaran"
            default_mapel = "Fisika"

        mapel = st.text_input(
            label_mapel,
            value=default_mapel,
            placeholder="Contoh: Fisika"
        )

        materi_pokok = st.text_input(
            "Fokus Topik / Materi Pokok",
            placeholder="Contoh: Usaha dan Energi"
        )

    with col2:

        alokasi_waktu = st.text_input(
            "Alokasi Waktu / Target JP",
            value=(
                "4 JP / Minggu"
                if is_sd
                else "3 JP / Minggu"
            )
        )

        profil_pancasila = st.multiselect(
            "Dimensi Profil Pelajar Pancasila",
            [
                "Beriman, Bertakwa kepada Tuhan YME, & Berakhlak Mulia",
                "Berkebinekaan Global",
                "Gotong Royong",
                "Mandiri",
                "Bernalar Kritis",
                "Kreatif",
            ],
            default=[
                "Bernalar Kritis",
                "Gotong Royong",
                "Mandiri",
            ],
        )

        daftar_22_perangkat = [
            "01. Analisis Capaian Pembelajaran (CP) & Pemetaan Elemen",
            "02. Alur Tujuan Pembelajaran (ATP) Lengkap",
            "03. Program Tahunan (PROTA)",
            "04. Program Semester (PROMES)",
            "05. Kriteria Ketercapaian Tujuan Pembelajaran (KKTP)",
            "06. Modul Ajar Lengkap (Format Baku Kurikulum Merdeka)",
            "07. Lembar Kerja Peserta Didik (LKPD) Berdiferensiasi",
            "08. Modul / Panduan Projek Profil Pancasila (P5)",
            "09. Jurnal Mengajar Harian & Agenda Guru Terstruktur",
            "10. Format & Kisi-kisi Asesmen Diagnostik",
            "11. Kisi-kisi & Instrumen Asesmen Formatif",
            "12. Kisi-kisi, Naskah Soal, & Kunci Jawaban Asesmen Sumatif",
            "13. Rubrik Penilaian Kinerja, Portofolio, serta Proyek",
            "14. Rekapitulasi Daftar Nilai Kurikulum Merdeka",
            "15. Buku Presensi / Lembar Absensi Siswa",
            "16. Format Program Pembelajaran Remedial & Pengayaan",
            "17. Distribusi Sumber Belajar & Buku Teks Pembelajaran",
            "18. Analisis Alokasi Waktu Efektif",
            "19. Format Analisis Kuantitatif Butir Soal Evaluasi",
            "20. Jurnal Sikap & Catatan Dimensi Profil Pelajar Pancasila",
            "21. Panduan Layanan Bimbingan & Konsultasi Akademik",
            "22. Format Laporan Evaluasi Diri Guru & Rencana Tindak Lanjut",
        ]

        jenis_perangkat = st.selectbox(
            "Pilih Dokumen yang Diterbitkan",
            daftar_22_perangkat
        )

    semester = st.radio(
        "Semester Berjalan",
        [
            "Semester Ganjil",
            "Semester Genap",
            "Semester Ganjil & Genap (1 Tahun Penuh)",
        ],
        horizontal=True,
    )


# =========================================================
# IDENTITAS SEKOLAH
# =========================================================

tab1, tab2 = st.tabs(
    [
        "🏛️ 2. Identitas Satuan Pendidikan",
        "✍️ 3. Pejabat & Penandatangan",
    ]
)


with tab1:

    c1, c2 = st.columns(2)

    with c1:

        dinas_pendidikan = st.text_input(
            "Dinas Pendidikan Pembina",
            value=(
                "DINAS PENDIDIKAN KOTA PONTIANAK"
                if is_sd
                else
                "DINAS PENDIDIKAN PROVINSI KALIMANTAN BARAT"
            )
        )

        nama_sekolah = st.text_input(
            "Nama Satuan Pendidikan",
            value=(
                "SDN 01 PONTIANAK"
                if is_sd
                else
                "SMAS NUSA HARAPAN"
            )
        )

    with c2:

        alamat_sekolah = st.text_input(
            "Alamat & Kontak Sekolah",
            value="Jl. Pancasila No. 10, Telp. (0561) 734567"
        )

        kota_sekolah = st.text_input(
            "Kota / Kabupaten Domisili",
            value="Pontianak"
        )


# =========================================================
# PEJABAT
# =========================================================

with tab2:

    c1, c2, c3 = st.columns(3)

    with c1:

        st.markdown(
            f"**{jabatan_guru_otomatis}**"
        )

        guru_nama = st.text_input(
            f"Nama & Gelar ({jabatan_guru_otomatis})",
            value=st.session_state.user_name,
            key="g_nama",
        )

        guru_nip = st.text_input(
            "NIP Guru",
            value="-",
            key="g_nip",
        )

    with c2:

        st.markdown(
            "**Kepala Sekolah**"
        )

        ks_nama = st.text_input(
            "Nama Kepala Sekolah",
            placeholder="Contoh: ZULKIFLI, S.Pd.",
            key="ks_nama",
        )

        ks_nip = st.text_input(
            "NIP Kepala Sekolah",
            placeholder="Contoh: 197508122005011004",
            key="ks_nip",
        )

    with c3:

        st.markdown(
            "**Pengawas Pembina**"
        )

        pengawas_nama = st.text_input(
            "Nama Pengawas Pembina",
            placeholder="Contoh: Andar",
            key="p_nama",
        )

        pengawas_nip = st.text_input(
            "NIP Pengawas",
            placeholder="Contoh: 196811231993032003",
            key="p_nip",
        )


# =========================================================
# TERBITKAN DOKUMEN
# =========================================================

if st.button(
    "✨ Terbitkan Dokumen Administrasi Resmi",
    use_container_width=True
):

    # VALIDASI
    if not mapel.strip():

        st.warning(
            "⚠️ Mata pelajaran wajib diisi."
        )

        st.stop()

    if (
        not dinas_pendidikan.strip()
        or not nama_sekolah.strip()
    ):

        st.warning(
            "⚠️ Data instansi belum lengkap."
        )

        st.stop()

    if not materi_pokok.strip():

        st.warning(
            "⚠️ Fokus topik/materi pokok wajib diisi."
        )

        st.stop()


    # DATA
    data_input = {

        "dinas": dinas_pendidikan,

        "sekolah": nama_sekolah,

        "alamat": alamat_sekolah,

        "kota_sekolah": kota_sekolah,

        "jabatan_guru": jabatan_guru_otomatis,

        "guru_nama": guru_nama,

        "guru_nip": guru_nip,

        "ks_nama": ks_nama or "-",

        "ks_nip": ks_nip or "-",

        "pengawas_nama": pengawas_nama or "-",

        "pengawas_nip": pengawas_nip or "-",

        "mapel": mapel,

        "materi_pokok": materi_pokok,

        "profil_pancasila": profil_pancasila,

        "fase_kelas": fase_kelas,

        "semester": semester,

        "alokasi_waktu": alokasi_waktu,

        "jenis_perangkat": jenis_perangkat,

        "tanggal_hari_ini": datetime.now().strftime(
            "%d-%m-%Y"
        ),
    }


    prompt_final = buat_instruksi_prompt(
        data_input
    )

    progress = st.progress(0)
    status = st.empty()


    # =====================================================
    # GEMINI
    # =====================================================

    try:

        status.info(
            "⚡ Menginisialisasi Google Gemini..."
        )

        client = genai.Client(
            api_key=st.session_state.user_api_key
        )

        progress.progress(15)


        # =================================================
        # MODEL YANG DIGUNAKAN
        # =================================================

        model_list = [
            "gemini-3.8-flash",
            "gemini-2.5-flash",
        ]


        response = None
        model_berhasil = None
        errors = []


        # =================================================
        # COBA MODEL
        # =================================================

        for model_name in model_list:

            if response:
                break

            for percobaan in range(3):

                try:

                    status.info(
                        f"📝 Menghasilkan dokumen dengan "
                        f"{model_name} "
                        f"(percobaan {percobaan + 1}/3)"
                    )

                    # PENTING:
                    # Tidak menggunakan temperature.
                    # Ini menghindari parameter yang tidak
                    # diperlukan pada model Gemini terbaru.

                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt_final,
                    )


                    if (
                        response
                        and getattr(
                            response,
                            "text",
                            None
                        )
                    ):

                        model_berhasil = model_name

                        break


                except Exception as exc:

                    error_text = str(exc)

                    errors.append(
                        f"{model_name} | "
                        f"Percobaan {percobaan + 1} | "
                        f"{error_text}"
                    )


                    # =====================================
                    # 503 / UNAVAILABLE
                    # =====================================

                    if (
                        "503" in error_text
                        or "UNAVAILABLE" in error_text.upper()
                    ):

                        if percobaan < 2:

                            waktu_tunggu = (
                                2 ** percobaan
                            )

                            status.warning(
                                f"⏳ Server Gemini sedang sibuk. "
                                f"Mencoba kembali dalam "
                                f"{waktu_tunggu} detik..."
                            )

                            time.sleep(
                                waktu_tunggu
                            )

                            continue


                    # =====================================
                    # ERROR LAIN
                    # =====================================

                    break


            if response:
                break


        progress.progress(90)


        # =================================================
        # SEMUA MODEL GAGAL
        # =================================================

        if (
            not response
            or not getattr(
                response,
                "text",
                None
            )
        ):

            detail_error = (
                "\n".join(errors[-6:])
                if errors
                else "Tidak ada respons dari Gemini."
            )

            raise RuntimeError(
                "Semua model Gemini gagal memproses "
                "permintaan.\n\n"
                + detail_error
            )


        # =================================================
        # SIMPAN HASIL
        # =================================================

        st.session_state.hasil_teks = (
            response.text
        )


        # =================================================
        # NAMA FILE
        # =================================================

        nama_perangkat_bersih = re.sub(
            r"^\d+\.\s*",
            "",
            jenis_perangkat
        )

        nama_file = (
            f"{nama_perangkat_bersih}_"
            f"{mapel}_"
            f"{fase_kelas}"
        )

        st.session_state.nama_file_base = re.sub(
            r"[^A-Za-z0-9_-]+",
            "_",
            nama_file
        ).strip("_")


        # =================================================
        # BERHASIL
        # =================================================

        progress.progress(100)

        status.success(
            f"✅ Dokumen berhasil diterbitkan "
            f"menggunakan {model_berhasil}."
        )

        time.sleep(0.8)

        status.empty()
        progress.empty()


    except Exception as exc:

        progress.empty()
        status.empty()

        st.error(
            "❌ Gagal menerbitkan dokumen."
        )

        st.code(
            str(exc),
            language="text"
        )


# =========================================================
# PREVIEW DAN DOWNLOAD
# =========================================================

if st.session_state.hasil_teks:

    st.markdown(
        "### 📄 Preview Lembar Kerja A4"
    )

    preview_html = markdown_to_html(
        st.session_state.hasil_teks
    )

    st.markdown(
        f"""
        <div class="paper-a4">
            {preview_html}
        </div>
        """,
        unsafe_allow_html=True,
    )


    c1, c2 = st.columns(2)


    with c1:

        docx_file = buat_file_docx(
            st.session_state.hasil_teks
        )

        st.download_button(
            "📄 Unduh Berkas Word (.DOCX)",
            data=docx_file,
            file_name=(
                f"{st.session_state.nama_file_base}.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )


    with c2:

        st.download_button(
            "📝 Unduh Teks Mentah (.TXT)",
            data=st.session_state.hasil_teks,
            file_name=(
                f"{st.session_state.nama_file_base}.txt"
            ),
            mime="text/plain",
            use_container_width=True,
        )


# =========================================================
# FOOTER
# =========================================================

st.markdown(
    """
    <div class="footer-box">
        SIAP AJAR 22 • Creator: Andar<br>
        © 2026 SIAP AJAR 22 • Engine AI Pembelajaran
    </div>
    """,
    unsafe_allow_html=True,
)
